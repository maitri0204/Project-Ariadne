from flask import Flask, render_template, request, jsonify, send_file
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import json
import re
import unicodedata
from typing import TypedDict, List, Dict, Any, Annotated
import operator
from langgraph.graph import StateGraph, END
from dotenv import load_dotenv

app = Flask(__name__)

# Load env explicitly (critical for Gunicorn)
load_dotenv("/var/www/portfolio_app/.env")

def get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set")
    return OpenAI(api_key=api_key)

def sanitize_filename(text):
    return re.sub(r'[^a-zA-Z0-9_-]', '_', text)

def build_skill_action_guidance(highest_skills, skill_percentages):
    """
    Converts selected highest skills with percentages into concrete guidance
    that prioritizes higher percentage skills in action plans.
    """
    skill_map = {
        "Strategy": "Prioritize planning, frameworks, option analysis, roadmaps, and decision memos. Provide clear goals, assumptions, risks, and trade-offs.",
        "Execution": "Prioritize weekly schedules, checklists, sprint plans, deliverables, accountability, and measurable outputs.",
        "Intellect": "Prioritize deep learning plans, conceptual clarity, problem-solving drills, and structured study methods.",
        "Asthetic": "Prioritize portfolio presentation, design taste, storytelling, and improvement loops via feedback and iterations.",
        "Balance": "Prioritize time-blocking, burnout prevention, sustainable habits, and realistic pacing with recovery routines.",
        "Movement": "Prioritize activity-based learning, practice-first tasks, daily reps, and performance routines.",
        "Expression": "Prioritize communication deliverables: presentations, writing, speaking practice, pitch decks, and content creation plans.",
        "Articulation": "Prioritize structured writing, clarity, argument building, and explanation skill. Add rubrics and review checklists.",
        "Observation": "Prioritize research, analysis, pattern recognition, journaling, case studies, and structured reflection.",
        "Ecological": "Prioritize systems thinking, context awareness, collaboration, stakeholder mapping, and environment/industry scanning."
    }
    
    selected = highest_skills or []
    percentages = skill_percentages or {}
    
    # Sort skills by percentage (highest first)
    sorted_skills = sorted(selected, key=lambda s: percentages.get(s, 0), reverse=True)
    
    lines = []
    lines.append("SKILL PRIORITIZATION (ordered by percentage - focus MORE on higher percentages):")
    
    for s in sorted_skills:
        percentage = percentages.get(s, 0)
        if s in skill_map:
            lines.append(f"- {s} ({percentage}%): {skill_map[s]}")
    
    if not lines or len(lines) == 1:
        return "- No dominant skill guidance available; still provide an actionable plan with timelines and measurable checkpoints."
    
    lines.append("\nIMPORTANT: Weight your recommendations based on these percentages. Higher percentage skills should get MORE action items and emphasis.")
    
    return "\n".join(lines)

def format_skills_with_percentages(skills, percentages):
    """
    Formats skills with their percentages for display in prompts.
    Example: "Strategy (14%), Observation (10%)"
    """
    if not skills:
        return 'NA'
    
    formatted = []
    for skill in skills:
        percentage = percentages.get(skill, 0)
        if percentage:
            formatted.append(f"{skill} ({percentage}%)")
        else:
            formatted.append(skill)
    
    return ', '.join(formatted)


@app.route('/')
def index():
    return render_template('index.html')

# =============================================================================
# MULTI-AGENT SYSTEM - STATE DEFINITION
# =============================================================================

class ReportState(TypedDict):
    """Shared state across all agents in the workflow"""
    report_type: str
    inputs: Dict[str, Any]
    sections_to_generate: List[str]
    current_section_index: int
    generated_sections: List[Dict[str, str]]  # ✅ REMOVED operator.add
    current_section_name: str
    current_section_content: str
    validation_result: Dict[str, Any]
    retry_count: int
    final_report: str
    error: str

# =============================================================================
# AGENT 1: SUPERVISOR AGENT
# =============================================================================

def supervisor_agent(state: ReportState) -> ReportState:
    """
    Supervisor Agent: Orchestrates the entire multi-agent workflow.
    Determines which sections need to be generated based on report type.
    """
    print(f"\n{'='*70}")
    print(f"[SUPERVISOR AGENT] Initializing {state['report_type'].upper()} Report Generation")
    print(f"{'='*70}\n")
    
    # Determine sections based on report type
    if state['report_type'] == 'career':
        sections = [
            "1. Detailed Career Role Breakdown",
            "2. Industry Specific Requirements",
            "3. Emerging Trends and Future Job Prospects",
            "4. Recommended Internships",
            "5. Professional Networking and Industry Associations",
            "6. Guidelines for Progress Monitoring & Support",
        ]
    else:  # development
        sections = [
            "1. Academic Interventions",
            "2. Non-Academic Interventions",
            "3. Habit Reengineering",
            "4. Physical Grooming",
            "5. Psychological Grooming",
            "6. Suggested Reading",
            "7. Health Discipline",
        ]
    
    state['sections_to_generate'] = sections
    state['current_section_index'] = 0
    state['generated_sections'] = []  # ✅ Initialize as empty list
    state['retry_count'] = 0
    state['error'] = ''
    
    print(f"[SUPERVISOR] Workflow Plan: {len(sections)} sections identified")
    for idx, section in enumerate(sections, 1):
        print(f"  {idx}. {section}")
    print()
    
    return state

# =============================================================================
# AGENT 2: SECTION GENERATOR AGENT
# =============================================================================

def section_generator_agent(state: ReportState) -> ReportState:
    """
    Section Generator Agent: Generates content for a specific section.
    Uses OpenAI GPT-4o with structured prompts and student profile data.
    """
    current_idx = state['current_section_index']
    section_name = state['sections_to_generate'][current_idx]
    
    print(f"[GENERATOR AGENT] Processing Section {current_idx + 1}/{len(state['sections_to_generate'])}: {section_name}")
    
    # Build the base prompt with student data
    skill_guidance = build_skill_action_guidance(
        state['inputs'].get("highest_skills", []),
        state['inputs'].get("skillpercentages", {})
    )
    
    base_prompt = (
        "Based on the following student profile information, you will write ONLY ONE "
        "of the requested sections of the report each time. "
        "Do not repeat the input data in the output.\n\n"
        "INPUT DATA:\n"
        f"- Highest Skills with Percentages: {format_skills_with_percentages(state['inputs'].get('highest_skills', []), state['inputs'].get('skillpercentages', {}))}\n"
        f"- Thinking Pattern: {state['inputs'].get('thinking_pattern', 'NA')}\n"
        f"- Achievement Style with Percentages: {format_skills_with_percentages(state['inputs'].get('achievement_style', []), state['inputs'].get('achievementpercentages', {}))}\n"
        f"- Learning Communication Style with Percentages: {format_skills_with_percentages(state['inputs'].get('learning_communication_style', []), state['inputs'].get('learningpercentages', {}))}\n"
        f"- Quotients with Percentages: {format_skills_with_percentages(state['inputs'].get('quotients', []), state['inputs'].get('quotientpercentages', {}))}\n"
        f"- Personality Type: {state['inputs'].get('personality_type', 'NA')}\n"
        f"- Suggested Career Roles: {state['inputs'].get('career_roles', 'NA')}\n\n"
        "CRITICAL INSTRUCTION - CAREER ROLE FOCUS:\n"
        f"YOU MUST USE THE EXACT CAREER ROLES ENTERED BY THE STUDENT: {state['inputs'].get('career_roles', 'NA')}\n"
        "DO NOT CHANGE OR SUGGEST DIFFERENT ROLES.\n"
        "The ENTIRE report must be centered around THESE EXACT ROLES ONLY.\n"
        "If multiple roles are mentioned (e.g., Software Engineer, Data Scientist, Event Manager), you MUST:\n"
        "  - Address ALL roles mentioned by the student\n"
        "  - Create separate subsections for each role where appropriate\n"
        "Every recommendation must be DIRECTLY relevant to the entered roles.\n"
        "Do NOT provide generic career advice. Do NOT suggest alternative roles.\n\n"
        f"{skill_guidance}\n\n"
        "NON-NEGOTIABLE OUTPUT RULES:\n"
        "- BE CONCISE AND DIRECT. NO lengthy explanations or verbose paragraphs.\n"
        "- Use BULLET POINTS for all lists and action items.\n"
        "- Each bullet should be 1-2 lines maximum.\n"
        "- NO introductory or concluding paragraphs.\n"
        "- Get straight to the actionable information.\n"
        "- TOTAL section length: 200-350 words maximum (not 500-700).\n"
        "- Do NOT use emojis or decorative symbols like * or #.\n"
        "- Format with clear subheadings for each role if multiple roles exist.\n\n"
        "FORMAT REQUIREMENTS:\n"
        "- First line must be the exact section heading.\n"
        "- Organize content with concise subheadings.\n"
        "- Use bullet points starting with '- ' for all lists.\n"
        "- Keep explanations minimal - focus on facts and action items.\n\n"
    )
    
    # Get section-specific prompt
    section_prompts = generate_section_prompts(state['report_type'], state['inputs'])
    section_prompt = section_prompts[current_idx]
    
    try:
        # Call OpenAI API
        client = get_openai_client()
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are academic and career expert having more than 35 years of experience "
                        "with deep knowledge in psychology, career development, and "
                        "personalized education planning. "
                        "You MUST strictly follow the career roles provided by the student. "
                        "Do NOT suggest different roles. Do NOT talk about generic or unspecified roles. "
                        "You must not use emojis or decorative symbols like * or # in the content."
                    ),
                },
                {"role": "user", "content": section_prompt},
            ],
            temperature=0.7,
            max_tokens=5500,
        )
        
        section_content = response.choices[0].message.content.strip()
        
        state['current_section_name'] = section_name
        state['current_section_content'] = section_content
        
        print(f"[GENERATOR AGENT] ✓ Content generated successfully ({len(section_content)} chars)")
        
    except Exception as e:
        print(f"[GENERATOR AGENT] ✗ Error: {str(e)}")
        state['error'] = f"Generation failed for {section_name}: {str(e)}"
        state['current_section_content'] = f"{section_name}\n\nContent generation failed."
    
    return state

# =============================================================================
# AGENT 3: VALIDATOR AGENT
# =============================================================================

def validator_agent(state: ReportState) -> ReportState:
    """
    Validator Agent: Validates generated content for quality and completeness.
    """
    section_name = state['current_section_name']
    section_content = state['current_section_content']
    
    print(f"[VALIDATOR AGENT] Analyzing: {section_name}")
    
    validation_result = {
        'is_valid': True,
        'issues': [],
        'requires_retry': False
    }
    
    # Special validation for Health Discipline
    if section_name == "7. Health Discipline":
        required_categories = ["Food", "Sleeping Discipline", "Hydration", "Lifestyle"]
        missing_categories = []
        
        for category in required_categories:
            if category.lower() not in section_content.lower():
                missing_categories.append(category)
        
        if missing_categories:
            validation_result['is_valid'] = False
            validation_result['issues'].append(f"Missing required categories: {', '.join(missing_categories)}")
            validation_result['requires_retry'] = True
            print(f"[VALIDATOR] ✗ FAILED: Missing categories - {missing_categories}")
        else:
            print(f"[VALIDATOR] ✓ All 4 categories present (Food, Sleep, Hydration, Lifestyle)")
    
    # General validation: Content length
    if len(section_content.strip()) < 100:
        validation_result['is_valid'] = False
        validation_result['issues'].append("Content too short (< 100 chars)")
        validation_result['requires_retry'] = True
        print(f"[VALIDATOR] ✗ Content too short")
    else:
        print(f"[VALIDATOR] ✓ Content length adequate")
    
    state['validation_result'] = validation_result
    
    if validation_result['is_valid']:
        print(f"[VALIDATOR] ✓✓ VALIDATION PASSED\n")
    else:
        print(f"[VALIDATOR] ✗✗ VALIDATION FAILED: {validation_result['issues']}\n")
    
    return state

# =============================================================================
# WORKFLOW CONTROL NODES
# =============================================================================

def should_retry_section(state: ReportState) -> str:
    """Decision node: Determine if section needs regeneration"""
    validation = state['validation_result']
    max_retries = 2
    
    if validation['requires_retry'] and state['retry_count'] < max_retries:
        state['retry_count'] += 1
        print(f"[DECISION NODE] Retry triggered (Attempt {state['retry_count']}/{max_retries})\n")
        return "retry"
    else:
        if state['retry_count'] >= max_retries and not validation['is_valid']:
            print(f"[DECISION NODE] ⚠ Max retries reached - accepting current version\n")
        else:
            print(f"[DECISION NODE] Section accepted\n")
        return "accept"

def save_section_and_continue(state: ReportState) -> ReportState:
    """Saves validated section and prepares for next section"""
    # ✅ CRITICAL FIX: Check if section already exists before appending
    section_already_exists = any(
        s['name'] == state['current_section_name'] 
        for s in state['generated_sections']
    )
    
    if not section_already_exists:
        state['generated_sections'].append({
            'name': state['current_section_name'],
            'content': state['current_section_content']
        })
        print(f"[WORKFLOW] ✓ Section stored: {state['current_section_name']}")
    else:
        print(f"[WORKFLOW] ⚠ Section '{state['current_section_name']}' already exists, skipping duplicate")
    
    print(f"[WORKFLOW] Progress: {len(state['generated_sections'])}/{len(state['sections_to_generate'])} sections saved\n")
    
    # Move to next section
    state['current_section_index'] = state['current_section_index'] + 1
    state['retry_count'] = 0
    
    return state

def has_more_sections(state: ReportState) -> str:
    """Decision node: Check if workflow should continue or finalize"""
    # ✅ SAFETY CHECK: Use generated_sections count as source of truth
    if len(state['generated_sections']) >= len(state['sections_to_generate']):
        print(f"[WORKFLOW] ✓ All {len(state['sections_to_generate'])} sections complete\n")
        return "finalize"
    
    if state['current_section_index'] < len(state['sections_to_generate']):
        next_section = state['sections_to_generate'][state['current_section_index']]
        print(f"[WORKFLOW] → Moving to next section: {next_section}\n")
        return "continue"
    else:
        print(f"[WORKFLOW] → All sections completed, finalizing report\n")
        return "finalize"

def finalize_report(state: ReportState) -> ReportState:
    """Final assembly: Combines all generated sections into complete report"""
    print(f"{'='*70}")
    print(f"[FINALIZER AGENT] Assembling final report")
    print(f"{'='*70}\n")
    
    all_content = []
    for idx, section in enumerate(state['generated_sections'], 1):
        all_content.append(section['content'])
        print(f"  ✓ Section {idx}: {section['name']}")
    
    state['final_report'] = "\n\n".join(all_content)
    
    print(f"\n[FINALIZER] ✓ Report assembly complete")
    print(f"[FINALIZER] Total length: {len(state['final_report'])} characters")
    print(f"[FINALIZER] Total sections: {len(state['generated_sections'])}")
    print(f"\n{'='*70}")
    print(f"[MULTI-AGENT WORKFLOW] Successfully Completed")
    print(f"{'='*70}\n")
    
    return state

# =============================================================================
# LANGGRAPH WORKFLOW BUILDER
# =============================================================================

def create_multi_agent_workflow():
    """
    Builds the complete multi-agent workflow using LangGraph.
    """
    workflow = StateGraph(ReportState)
    
    # Register all agent nodes
    workflow.add_node("supervisor", supervisor_agent)
    workflow.add_node("generator", section_generator_agent)
    workflow.add_node("validator", validator_agent)
    workflow.add_node("save_and_continue", save_section_and_continue)
    workflow.add_node("finalizer", finalize_report)
    
    # Define workflow edges
    workflow.set_entry_point("supervisor")
    
    workflow.add_edge("supervisor", "generator")
    workflow.add_edge("generator", "validator")
    
    workflow.add_conditional_edges(
        "validator",
        should_retry_section,
        {
            "retry": "generator",
            "accept": "save_and_continue"
        }
    )
    
    workflow.add_conditional_edges(
        "save_and_continue",
        has_more_sections,
        {
            "continue": "generator",
            "finalize": "finalizer"
        }
    )
    
    workflow.add_edge("finalizer", END)
    
    return workflow.compile()

# =============================================================================
# MAIN ENTRY POINT FOR MULTI-AGENT REPORT GENERATION
# =============================================================================

def generate_report_with_agents(report_type: str, inputs: Dict[str, Any]) -> str:
    """
    Main function to generate report using multi-agent system.
    """
    workflow = create_multi_agent_workflow()
    
    initial_state = {
        'report_type': report_type,
        'inputs': inputs,
        'sections_to_generate': [],
        'current_section_index': 0,
        'generated_sections': [], 
        'current_section_name': '',
        'current_section_content': '',
        'validation_result': {},
        'retry_count': 0,
        'final_report': '',
        'error': ''
    }
    
    final_state = workflow.invoke(initial_state)
    
    return final_state['final_report']

def generate_section_prompts(report_type, inputs):
    skill_guidance = build_skill_action_guidance(
        inputs.get("highest_skills", []), 
        inputs.get("skillpercentages", {})
    )

    base_prompt = (
        "Based on the following student profile information, you will write ONLY ONE "
        "of the requested sections of the report each time. "
        "Do not repeat the input data in the output.\n\n"
        "INPUT DATA:\n"
        f"- Standard / Year: {inputs.get('standard', 'NA')}\n"
        f"- Board: {inputs.get('board', 'NA')}\n"
        f"- Highest Skills with Percentages: {format_skills_with_percentages(inputs.get('highest_skills', []), inputs.get('skillpercentages', {}))}\n"
        f"- Thinking Pattern: {inputs.get('thinking_pattern', 'NA')}\n"
        f"- Achievement Style with Percentages: {format_skills_with_percentages(inputs.get('achievement_style', []), inputs.get('achievementpercentages', {}))}\n"
        f"- Learning Communication Style with Percentages: {format_skills_with_percentages(inputs.get('learning_communication_style', []), inputs.get('learningpercentages', {}))}\n"
        f"- Quotients with Percentages: {format_skills_with_percentages(inputs.get('quotients', []), inputs.get('quotientpercentages', {}))}\n"
        f"- Personality Type: {inputs.get('personality_type', 'NA')}\n"
        f"- Suggested Career Roles: {inputs.get('career_roles', 'NA')}\n\n"
        "CRITICAL AGE/LEVEL ADAPTATION RULE:\n"
        "You MUST adapt all advice to the student's Standard/Year.\n"
        "- If the student is in school (e.g., 6th-12th): focus on school-level actions, subject foundations, study routines, age-appropriate internships/projects (mini-projects), and parent/teacher support.\n"
        "- If the student is in college (e.g., FY/SY/TY/1st-4th year): focus on industry readiness, internships, projects, networking, resume/portfolio, and placement preparation.\n"
        "Do NOT give college-level internship/placement advice to an 8th/9th student.\n"
        "Do NOT give school timetable advice to a final-year college student.\n\n"
        "CRITICAL INSTRUCTION - CAREER ROLE FOCUS:\n"
        f"YOU MUST USE THE EXACT CAREER ROLES ENTERED BY THE STUDENT: {inputs.get('career_roles', 'NA')}\n"
        "DO NOT CHANGE OR SUGGEST DIFFERENT ROLES.\n"
        "The ENTIRE report must be centered around THESE EXACT ROLES ONLY.\n"
        "If multiple roles are mentioned (e.g., Software Engineer, Data Scientist, Event Manager), you MUST:\n"
        "  - Address ALL roles mentioned by the student\n"
        "  - Create separate subsections for each role where appropriate\n"
        "Every recommendation must be DIRECTLY relevant to the entered roles.\n"
        "Do NOT provide generic career advice. Do NOT suggest alternative roles.\n\n"
        f"{skill_guidance}\n\n"
        "BOARD CONTEXT RULE:\n"
        "You MUST adapt learning methods, practice strategies, and execution “commands” (how skills are practiced, reinforced, and assessed) according to the student’s academic Board and study environment.\n"
        "CBSE (Central Board of Secondary Education): emphasize NCERT-based conceptual clarity, structured syllabus progression, exam-oriented practice, problem-solving aligned with national competitive exams, and strong fundamentals with time-bound revision.\n"
        "ICSE (Indian Certificate of Secondary Education): emphasize detailed conceptual understanding, strong language and writing skills, descriptive answer structuring, curriculum depth, and analytical explanation alongside exam preparedness.\n"
        "State Board (e.g., Gujarat Board, UP Board, etc.): emphasize board-pattern questions, language-medium sensitivity where applicable, strengthening core fundamentals, confidence-building through guided practice, and deliberate bridging of gaps toward national-level competitive and industry standards.\n"
        "IB / Cambridge: emphasize inquiry-based learning, conceptual depth, project and research work, independent thinking, report writing, interdisciplinary understanding, and application-oriented assessments.\n"
        "Learning activities, practice intensity, expectations, and support mechanisms MUST be realistically aligned with the academic exposure, assessment style, and learning culture of the student’s Board.\n"
        "Keep recommendations practical and appropriate to the board.\n\n"
        "NON-NEGOTIABLE OUTPUT RULES:\n"
        "- BE CONCISE AND DIRECT. NO lengthy explanations or verbose paragraphs.\n"
        "- Use BULLET POINTS for all lists and action items.\n"
        "- Each bullet should be 1-2 lines maximum.\n"
        "- NO introductory or concluding paragraphs.\n"
        "- Get straight to the actionable information.\n"
        "- TOTAL section length: 200-350 words maximum (not 500-700).\n"
        "- Do NOT use emojis or decorative symbols like * or #.\n"
        "- Format with clear subheadings for each role if multiple roles exist.\n\n"
        "FORMAT REQUIREMENTS:\n"
        "- First line must be the exact section heading.\n"
        "- Organize content with concise subheadings.\n"
        "- Use bullet points starting with '- ' for all lists.\n"
        "- Keep explanations minimal - focus on facts and action items.\n\n"
    )

    if report_type == "career":
        sections = [
            "1. Detailed Career Role Breakdown",
            "2. Industry Specific Requirements",
            "3. Emerging Trends and Future Job Prospects",
            "4. Recommended Internships",
            "5. Professional Networking and Industry Associations",
            "6. Guidelines for Progress Monitoring & Support",
        ]

    else:  # development
        sections = [
            "1. Academic Interventions",
            "2. Non-Academic Interventions",
            "3. Habit Reengineering",
            "4. Physical Grooming",
            "5. Psychological Grooming",
            "6. Suggested Reading",
            "7. Health Discipline",
        ]

    # Section-specific blueprints for more targeted outputs
    section_blueprints = {
        "1. Detailed Career Role Breakdown": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\n"
            "\n"
            "For EACH entered career role, output in this EXACT format (one role per block):\n"
            "\n"
            "Career Role: [Role Name]\n"
            "Technical Skills: [comma-separated list]\n"
            "Soft Skills: [comma-separated list]\n"
            "Undergraduate Education: [degree name]\n"
            "Postgraduate Education: [degree name]\n"
            "Micro-degrees: [comma-separated certifications]\n"
            "Certifications: [comma-separated list]\n"
            "Career Progression: [progression path with arrows]\n"
            "Salary Range: [amount and currency]\n"
            "Day-to-Day Responsibilities: [comma-separated list]\n"
            "\n"
            "[Leave blank line between roles]\n"
            "\n"
            "DO NOT CREATE MARKDOWN TABLES. DO NOT USE PIPES |\n"
            "Each field on its own line with label: value format.\n"
        ),
        "2. Industry Specific Requirements": (
            "SECTION-SPECIFIC REQUIREMENTS:\n"
            "- For EACH career role, organize requirements in BEGINNER → ADVANCED progression\n"
            "- Create a TABLE with columns: Level | Certification Name | Application Process | Duration | Assistance Resources\n"
            "- Structure for each role:\n"
            "  **For [Career Role Name]:**\n"
            "  \n"
            "  Beginner Level:\n"
            "  - Certification Name: [Name]\n"
            "  - Application Process: Step-by-step how to apply (registration website, prerequisites, exam format)\n"
            "  - Duration: Time to complete (e.g., 3 months, 6 weeks)\n"
            "  - Assistance Resources: Where to get help (official courses, study materials, forums, coaching)\n"
            "  \n"
            "  Intermediate Level:\n"
            "  [same format]\n"
            "  \n"
            "  Advanced Level:\n"
            "  [same format]\n"
            "- Include ALL important details: registration links, prerequisites, exam format, study resources, cost (if applicable)\n"
            "- Be SPECIFIC and ACTIONABLE - students should be able to act on this information immediately\n"
        ),
        "3. Emerging Trends and Future Job Prospects": (
            "SECTION-SPECIFIC REQUIREMENTS:\n"
            "- Determine the CURRENT YEAR dynamically at the time of report generation.\n"
            "- Define time ranges as follows:\n"
            "  * Past Trend: Previous 3 completed years (Current Year - 3 to Current Year - 1)\n"
            "  * Present Trend: Current Year\n"
            "  * Future Prediction: Next 3 years (Current Year + 1 to Current Year + 3)\n"
            "\n"
            "- For EACH career role, create a separate subsection with clear heading:\n"
            "  **[Career Role Name]**\n"
            "- Then provide a TABLE with columns:\n"
            "  Past Trend (Previous 3 Years) | Present Trend (Current Year) | Future Prediction (Next 3 Years)\n"
            "\n"
            "- Include STATISTICAL DATA based on real industry trends such as:\n"
            "  market size, job growth %, salary trends, technology adoption rates\n"
            "\n"
            "- Rows should cover:\n"
            "  * Job Demand Growth\n"
            "  * Average Salary Trends\n"
            "  * Key Technologies / Skills\n"
            "  * Industry Adoption Rate\n"
            "  * Geographic Demand\n"
            "\n"
            "- IMPORTANT:\n"
            "  * Use realistic, conservative estimates aligned with reputable industry reports.\n"
            "  * If exact figures are unavailable, provide clearly stated approximate ranges.\n"
            "  * DO NOT fabricate precise statistics or cite fake reports.\n"
        ),
        "4. Recommended Internships": (
            "SECTION-SPECIFIC REQUIREMENTS:\n"
            "- Organize by CAREER ROLE with clear role headings\n"
            "- For each role, provide a TABLE with columns: Internship Type | Industries (Small/Medium/Large) | Expected Outcomes\n"
            "- Structure:\n"
            "  **For [Career Role Name]:**\n"
            "  \n"
            "  Table with:\n"
            "  - Internship Type: Specific internship position (e.g., 'Data Analysis Intern', 'ML Engineering Intern')\n"
            "  - Industries: List industries across different scales:\n"
            "    * Small: Startups, boutique firms (mention 2-3 types)\n"
            "    * Medium: Mid-sized companies, regional firms (mention 2-3 types)\n"
            "    * Large: Fortune 500, multinational corporations (mention 2-3 types)\n"
            "  - Expected Outcomes: 3-4 key learning outcomes from that internship type\n"
            "- DO NOT use 'Point 1', 'Point 2' - use meaningful internship type names\n"
            "- Provide 5-8 internship types per role\n"
            "- Include application pipeline advice at the end (application strategy, platforms, timing)\n"
        ),
        "5. Professional Networking and Industry Associations": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\n"
            "\n"
            "For EACH entered career role, output in this EXACT format:\n"
            "\n"
            "For [Career Role Name]:\n"
            "\n"
            "Professional Associations:\n"
            "- [Association 1]\n"
            "- [Association 2]\n"
            "- [Association 3]\n"
            "- [Association 4]\n"
            "- [Association 5]\n"
            "\n"
            "Industry Events:\n"
            "- [Event/Conference 1]\n"
            "- [Event/Conference 2]\n"
            "- [Event/Conference 3]\n"
            "- [Event/Conference 4]\n"
            "- [Event/Conference 5]\n"
            "\n"
            "Networking Strategy:\n"
            "- [Strategy 1]\n"
            "- [Strategy 2]\n"
            "- [Strategy 3]\n"
            "- [Strategy 4]\n"
            "- [Strategy 5]\n"
            "\n"
            "[Leave blank line between roles]\n"
            "\n"
            "DO NOT CREATE MARKDOWN TABLES. DO NOT USE PIPES |\n"
            "Use bullet points (- ) for each item.\n"
        ),
        "6. Guidelines for Progress Monitoring & Support": (
            "SECTION-SPECIFIC REQUIREMENTS:\n"
            "- Create a HORIZONTAL comparison table with this structure:\n"
            "  | Aspect | [Career Role 1] | [Career Role 2] | [Career Role 3] |\n"
            "- Rows (Aspects) should include:\n"
            "  * Strategy (how to develop strategic skills for this role)\n"
            "  * Observation (practice exercises for observation skills)\n"
            "  * Balance (work-life balance techniques)\n"
            "  * Intellect (learning and problem-solving approaches)\n"
            "  * Expression (communication skill development)\n"
            "  * Execution (project delivery methods)\n"
            "  * KPIs (key performance indicators to track)\n"
            "  * Mentorship (how to find mentors)\n"
            "  * Self-Assessment (monthly review checklist)\n"
            "  * Feedback Loops (peer review, mock interviews, portfolio reviews)\n"
            "- This format allows EASY COMPARISON across all career roles\n"
            "- Keep each cell concise but actionable (2-3 sentences max)\n"
        ),
        # Development report sections remain the same
        "1. Academic Interventions": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\\n"
            "\\n"
            "Create a 3-year academic intervention plan WITH CONTENT FOR EVERY SINGLE MONTH (all 12 months each year).\\n"
            "You MUST output in this EXACT structure so it can be converted to Word tables:\\n"
            "\\n"
            "Year 1:\\n"
            "- Month: January\\n"
            " Activity: [what needs to be done]\\n"
            " Technical Skills: [comma-separated skills]\\n"
            " Soft Skills: [comma-separated skills]\\n"
            " Learning Material: [courses, books, platforms]\\n"
            " Objective: [1–2 line objective]\\n"
            "\\n"
            "- Month: February\\n"
            " Activity: [...]\\n"
            " Technical Skills: [...]\\n"
            " Soft Skills: [...]\\n"
            " Learning Material: [...]\\n"
            " Objective: [...]\\n"
            "\\n"
            "- Month: March\\n"
            "- Month: April\\n"
            "- Month: May\\n"
            "- Month: June\\n"
            "- Month: July\\n"
            "- Month: August\\n"
            "- Month: September\\n"
            "- Month: October\\n"
            "- Month: November\\n"
            "- Month: December\\n"
            "\\n"

            "[For each month: Activity, Technical Skills, Soft Skills, Learning Material, Objective]\\n"
            "\\n"
            "Year 2:\\n"
            "[Repeat EXACT same month-by-month structure for ALL 12 months (January-December) with different content]\\n"
            "\\n"
            "Year 3:\\n"
            "[Repeat EXACT same month-by-month structure for ALL 12 months (January-December) with different content]\\n"
            "\\n"

            "CRITICAL RULES:\\n"
            "- MUST include ALL 12 months for each year (January through December).\\n"
            "- Each month MUST have: Activity, Technical Skills, Soft Skills, Learning Material, Objective.\\n"
            "- DO NOT use markdown tables.\\n"
            "- DO NOT use pipes |.\\n"
            "- Use only the labels: Month, Activity, Technical Skills, Soft Skills, Learning Material, Objective.\\n"
            "- Ensure indentation exactly as shown (- Month, then 2-space indented fields).\\n"
            "- Make each month's content UNIQUE and PROGRESSIVE through the year.\\n"
        ),

        "2. Non-Academic Interventions": (
            "SECTION-SPECIFIC REQUIREMENTS – OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\n"
            "\n"
            "Create a COMPREHENSIVE 3-YEAR NON-ACADEMIC INTERVENTION PLAN focused on PERSONALITY DEVELOPMENT, LIFE SKILLS, SOCIAL INTELLIGENCE, EMOTIONAL MATURITY, DISCIPLINE, ETHICS, HEALTH, AND REAL-WORLD ADAPTABILITY.\n"
            "\n"
            "These interventions MUST NOT be academic courses, degrees, or syllabus-based learning. They must focus on experiential learning, behavioral conditioning, exposure-based growth, habit formation, emotional regulation, leadership readiness, and practical life competence.\n"
            "\n"
            "The plan must show CLEAR PROGRESSION across 3 years:\n"
            "- Year 1: Foundation building, self-awareness, discipline, exposure, basic social and life skills\n"
            "- Year 2: Skill strengthening, responsibility, leadership exposure, stress handling, independence\n"
            "- Year 3: Maturity, strategic thinking, resilience, ethical grounding, real-world readiness\n"
            "\n"
            "You MUST output in the EXACT structure below so it can be directly converted into Word tables. DO NOT change labels, order, or wording of fields.\n"
            "\n"

            "Year 1:\\n"
            "- Month: January\\n"
            " Activity: [clearly defined non-academic activity focused on behavior, exposure, or life skill development]\\n"
            " Technical Skills: [practical real-world skills such as organization, planning, observation, coordination, basic tools, systems thinking – comma-separated]\\n"
            " Soft Skills: [behavioral and psychological skills such as discipline, confidence, empathy, adaptability, communication – comma-separated]\\n"
            " Learning Outcome: [specific capability, behavior change, or internal skill the student will develop]\\n"
            " Objective: [1–2 lines explaining WHY this activity is included and what developmental gap it addresses]\\n"
            "\\n"
            "- Month: February\\n"
            "- Month: March\\n"
            "- Month: April\\n"
            "- Month: May\\n"
            "- Month: June\\n"
            "- Month: July\\n"
            "- Month: August\\n"
            "- Month: September\\n"
            "- Month: October\\n"
            "- Month: November\\n"
            "- Month: December\\n"
            "\\n"
            "[For EVERY month, you MUST provide ALL of the following: Activity, Technical Skills, Soft Skills, Learning Outcome, Objective. No field can be skipped.]\\n"
            "\\n"

            "Year 2:\\n"
            "Repeat the EXACT SAME STRUCTURE as Year 1 with ALL 12 months (January–December).\\n"
            "Year 2 activities must be MORE DEMANDING than Year 1 and focus on responsibility, leadership exposure, social confidence, stress tolerance, decision-making, and independence.\\n"
            "\\n"

            "Year 3:\\n"
            "Repeat the EXACT SAME STRUCTURE as Year 1 with ALL 12 months (January–December).\\n"
            "Year 3 activities must reflect MATURITY and REAL-WORLD READINESS, including leadership ownership, ethical judgment, resilience under pressure, strategic thinking, and long-term self-management.\\n"
            "\\n"

            "CRITICAL RULES (NON-NEGOTIABLE):\\n"
            "- ALL 3 years MUST include ALL 12 months from January to December.\\n"
            "- EACH MONTH MUST include ALL FIVE fields: Activity, Technical Skills, Soft Skills, Learning Outcome, Objective.\\n"
            "- Use the field name EXACTLY as 'Learning Outcome' (do NOT use learning material, resources, or books).\\n"
            "- NO academic subjects, exams, degrees, certifications, or classroom-style learning.\\n"
            "- Content must be NON-REPETITIVE, LOGICALLY PROGRESSIVE, and DEVELOPMENTALLY COHERENT across months and years.\\n"
            "- Activities must clearly contribute to emotional maturity, discipline, social competence, self-awareness, resilience, leadership, health, ethics, and life preparedness.\\n"
            "- Output must be plain text only, no markdown tables, no symbols, no pipes.\\n"
        ),

        "3. Habit Reengineering": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\n"
            "\n"
            "Design a structured 3-year habit reengineering plan focused on long-term student development.\n"
            "The plan should gradually build consistency, discipline, self-regulation, learning habits, and responsibility using small, repeatable actions rather than motivation.\n"
            "Each year must show clear progression from basic routine formation to advanced self-management and independent execution.\n"
            "\n"
            "You MUST output content for AT LEAST 6-7 months per year, spread across the year (not consecutive months).\n"
            "Suggested months to include: January, March, June, September, November, December (+ one additional month of your choice).\n"
            "\n"
            "Each month must include ONE primary habit-building focus aligned with academic discipline, personal responsibility, emotional regulation, or learning efficiency.\n"
            "Activities should be realistic, age-appropriate, and designed to create sustainable daily or weekly habits.\n"
            "\n"
            "You MUST output in this EXACT structure:\n"
            "\n"

            "Year 1:\n"
            "- Month: January\n"
            " Activity: [specific habit-building activity or routine]\n"
            " Action plan: [provide detailed steps to perform the activity]"
            " Objective: [clear purpose of this habit in 1–2 lines]\n"
            " Habits to Develop: [comma-separated daily or weekly habits]\n"
            " Soft Skills: [comma-separated behavioral or personal skills]\n"
            " Learning Outcomes: [observable outcomes or behavioral improvements]\n"
            "\n"
            "- Month: March\n"
            "- Month: June\n"
            "- Month: September\n"
            "- Month: November\n"
            "- Month: December\n"
            "\n"
            "Year 2:\n"
            "Repeat the same structure with AT LEAST 6-7 months spread across the year.\n"
            "Content must reflect higher responsibility, improved consistency, better time management, and increased self-awareness compared to Year 1.\n"
            "\n"
            "Year 3:\n"
            "Repeat the same structure with AT LEAST 6-7 months spread across the year.\n"
            "Content must focus on autonomy, long-term planning, self-discipline without supervision, and preparation for academic or career transitions.\n"
            "\n"

            "RULES:\n"
            "- Include at least 6-7 months per year (NOT all 12), spread throughout the year.\n"
            "- Use EXACT field names: Month, Activity, Action Plan, Objective, Habits to Develop, Soft Skills, Learning Outcomes.\n"
            "- No markdown tables, no pipes, no bullet nesting.\n"
            "- Each month's content must be UNIQUE, practical, and PROGRESSIVE across years.\n"
        ),

        "4. Physical Grooming": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\n"
            "\n"
            "Create a 3-year PHYSICAL GROOMING PLAN focused on HEALTH, DISCIPLINE, ENERGY MANAGEMENT, POSTURE, PROFESSIONAL PRESENCE, AND STRESS REGULATION.\n"
            "\n"
            "Physical Grooming must be treated as a DEVELOPMENTAL FOUNDATION that supports mental clarity, confidence, consistency, and long-term career readiness — not as fitness training or fashion alone.\n"
            "\n"
            "Activities should address: daily physical discipline, posture and body awareness, hygiene and self-care routines, nutrition and sleep regulation, physical confidence, stress reduction, and professional appearance readiness.\n"
            "\n"
            "You MUST output content for AT LEAST 6-7 months per year, spread across the year (NOT consecutive months), to reflect phased and sustainable physical development.\n"
            "Suggested months to include: January, April, June, September, October, December.\n"
            "\n"
            "You MUST output in this EXACT structure:\n"
            "\n"
            "Year 1:\n"
            "- Month: January\n"
            " Activity: [physical grooming activity focused on body awareness, routine formation, or basic health discipline]\n"
            " Objective: [1–2 lines explaining how this activity builds physical discipline, energy, confidence, or readiness]\n"
            " Physical & Mental Skills Developed: [comma-separated skills such as stamina, posture, focus, balance, stress control]\n"
            " Soft Skills: [comma-separated skills such as self-discipline, confidence, consistency, self-awareness]\n"
            " Learning Outcomes: [clear outcomes related to physical stability, mental clarity, and personal presentation]\n"
            "\n"
            "- Month: April\n"
            "- Month: June\n"
            "- Month: September\n"
            "- Month: October\n"
            "- Month: December\n"
            "\n"
            "Year 2:\n"
            "Repeat the SAME structure with AT LEAST 6-7 months spread throughout the year.\n"
            "Year 2 activities must show PROGRESSION toward improved stamina, posture, stress tolerance, hygiene discipline, and professional appearance.\n"
            "\n"
            "Year 3:\n"
            "Repeat the SAME structure with AT LEAST 6-7 months spread throughout the year.\n"
            "Year 3 activities must reflect MATURITY, SELF-MANAGEMENT, LEADERSHIP PRESENCE, AND LONG-TERM PHYSICAL SUSTAINABILITY.\n"
            "\n"
            "RULES:\n"
            "- Include at least 6-7 months per year, spread across the year (NOT all 12 months).\n"
            "- Use EXACT field names: Month, Activity, Objective, Physical & Mental Skills Developed, Soft Skills, Learning Outcomes.\n"
            "- No markdown tables, no pipes.\n"
            "- Each month's content must be UNIQUE, PURPOSEFUL, and DEVELOPMENTALLY PROGRESSIVE across the 3 years.\n"
            ),

        "5. Psychological Grooming": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\n"
            "\n"
            "Create a 3-year PSYCHOLOGICAL GROOMING PLAN focused on EMOTIONAL REGULATION, MENTAL CLARITY, STRESS MANAGEMENT, RESILIENCE, DECISION-MAKING, AND SELF-DISCIPLINE.\n"
            "\n"
            "Psychological Grooming must support sustained academic and career performance by strengthening emotional stability, pressure tolerance, motivation continuity, self-awareness, and reflective thinking.\n"
            "\n"
            "Activities should address: emotional awareness and control, stress response management, cognitive clarity, failure handling, motivation sustainability, confidence stabilization, and self-reflection.\n"
            "\n"
            "You MUST output content for AT LEAST 6-7 months per year, spread across the year (NOT consecutive months), to allow gradual and sustainable psychological development.\n"
            "Suggested months to include: January, February, June, August, November, December.\n"
            "\n"
            "You MUST output in this EXACT structure:\n"
            "\n"

            "Year 1:\\n"
            "- Month: January\\n"
            " Activity: [psychological grooming activity focused on self-awareness, emotional regulation, or mental discipline]\\n"
            " Objective: [1–2 lines explaining how this activity improves mental stability, focus, or emotional control]\\n"
            " Psychological Skills Developed: [comma-separated skills such as emotional regulation, focus, resilience, stress tolerance]\\n"
            " Soft Skills: [comma-separated skills such as self-discipline, confidence, adaptability, responsibility]\\n"
            " Learning Outcomes: [clear outcomes related to emotional stability, mental clarity, and behavioral control]\\n"
            "\\n"
            "- Month: February\\n"
            "- Month: June\\n"
            "- Month: August\\n"
            "- Month: November\\n"
            "- Month: December\\n"
            "\\n"
            "Year 2:\\n"
            "Repeat the SAME structure with AT LEAST 6-7 months spread throughout the year.\\n"
            "Year 2 activities must show PROGRESSION toward stress resilience, decision-making maturity, motivation stability, and pressure handling.\\n"
            "\\n"

            "Year 3:\\n"
            "Repeat the SAME structure with AT LEAST 6-7 months spread throughout the year.\\n"
            "Year 3 activities must reflect PSYCHOLOGICAL MATURITY, SELF-REGULATION, RESPONSIBILITY OWNERSHIP, AND LONG-TERM MENTAL ENDURANCE.\\n"
            "\\n"

            "RULES:\\n"
            "- Include at least 6-7 months per year, spread across the year (NOT all 12 months).\\n"
            "- Use EXACT field names: Month, Activity, Objective, Psychological Skills Developed, Soft Skills, Learning Outcomes.\\n"
            "- No markdown tables, no pipes.\\n"
            "- Each month's content must be UNIQUE, PURPOSEFUL, and PROGRESSIVELY BUILD psychological strength across the 3 years.\\n"
        ),

        "6. Suggested Reading": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\\n"
            "\\n"
            "You MUST output AT LEAST 15 books (minimum 15, preferably 18-20).\\n"
            "\\n"
            "CRITICAL BOOK SELECTION RULES:\\n"
            "- ALL books MUST be AVAILABLE IN INDIA (physically or as e-books on Amazon India, Flipkart, or popular Indian bookstores).\\n"
            "- DO NOT suggest books that are out of print, region-locked, or unavailable in India.\\n"
            "- DO NOT HALLUCINATE or make up book titles. ONLY suggest REAL, VERIFIED, FAMOUS books.\\n"
            "- Books should be a BALANCED MIX of:\\n"
            "  * TECHNICAL/DOMAIN BOOKS (50-60%): Directly related to the career role (e.g., finance, programming, data science, management).\\n"
            "  * SOFT SKILLS BOOKS (40-50%): Communication, leadership, emotional intelligence, productivity, mindset, time management, professional development.\\n"
            "\\n"
            "OUTPUT FORMAT (row-wise blocks with these EXACT fields):\\n"
            "\\n"
            "- Book Name: [title]\\n"
            "  Author: [author name]\\n"
            "  Publication: [publisher or edition]\\n"
            "  Availability in India: [Mention 'Available on Amazon India/Flipkart/Meesho' or specific Indian publisher]\\n"
            "  Why Should This Book Be Read?: [1–2 lines explaining relevance to their career AND skill development]\\n"
            "\\n"
            "[Repeat the above block for EACH BOOK - minimum 15 books, aim for 18-20]\\n"
            "\\n"
            "ORGANIZATION:\\n"
            "- Organize books by categories:\\n"
            "  **TECHNICAL/DOMAIN BOOKS** (8-10 books)\\n"
            "  **SOFT SKILLS & PROFESSIONAL DEVELOPMENT BOOKS** (7-10 books)\\n"
            "\\n"
            "CRITICAL RULES:\\n"
            "- MINIMUM 15 books. Aim for 18-20 books.\\n"
            "- Each book MUST include all 5 fields: Book Name, Author, Publication, Availability in India, Why Should This Book Be Read?.\\n"
            "- Make 'Why Should This Book Be Read?' specific to their career/skills (not generic).\\n"
            "- VERIFY that books are famous, well-reviewed, and actually available in India.\\n"
            "- Include ISBN or edition details if helpful for verification.\\n"
            "- No markdown tables, no pipes.\\n"
        ),

        "7. Health Discipline": (
            "SECTION-SPECIFIC REQUIREMENTS - OUTPUT STRUCTURED FORMAT (NOT MARKDOWN TABLE):\\n"
            "\\n"
            "CRITICAL: You MUST provide recommendations for ALL FOUR categories in this EXACT order:\\n"
            "1. FOOD (6-8 recommendations)\\n"
            "2. SLEEPING DISCIPLINE (5-6 recommendations)\\n"
            "3. HYDRATION (4-5 recommendations)\\n"
            "4. LIFESTYLE (5-6 recommendations)\\n"
            "\\n"
            "Recommendations across ALL 4 categories. DO NOT skip any category.\\n"
            "\\n"
            "===== CATEGORY 1: FOOD =====\\n"
            "Provide 6-8 specific food recommendations with these sub-categories:\\n"
            "- Balanced Diet with Whole Foods\\n"
            "- Morning: Warm Lemon Water & Soaked Nuts\\n"
            "- Breakfast: Protein-Rich Meal (Besan Chilla, Paneer Paratha)\\n"
            "- Mid-Morning Snack: Fruits (Banana, Apple, Orange, Papaya, Dry Fruits)\\n"
            "- Lunch: Dal, Roti, Rice, Green Vegetables, Salad, Curd\\n"
            "- Evening Snack: Herbal Tea & Roasted Makhana or Nuts\\n"
            "- Dinner: Light Meal (Khichdi, Vegetable Soup, Multigrain Roti with Sabzi)\\n"
            "- Bedtime: Warm Milk with Turmeric or Ashwagandha\\n"
            "\\n"
            "For each, output:\\n"
            "- Category: Food\\n"
            "  Recommendation: [specific food/meal]\\n"
            "  Benefits for Mental Health: [1-2 lines]\\n"
            "  Benefits for Physical Health: [1-2 lines]\\n"
            "\\n"
            "===== CATEGORY 2: SLEEPING DISCIPLINE =====\\n"
            "Provide 5-6 specific sleep recommendations:\\n"
            "- Maintaining a Fixed Sleep Schedule (10 PM - 6 AM)\\n"
            "- Avoiding Screens 1 Hour Before Bed\\n"
            "- Practicing Nighttime Meditation/Deep Breathing\\n"
            "- Using Dim Lights Before Sleeping\\n"
            "- Avoiding Heavy or Spicy Meals\\n"
            "\\n"
            "For each, output:\\n"
            "- Category: Sleeping Discipline\\n"
            "  Recommendation: [specific practice]\\n"
            "  Benefits for Mental Health: [1-2 lines]\\n"
            "  Benefits for Physical Health: [1-2 lines]\\n"
            "\\n"
            "===== CATEGORY 3: HYDRATION ===== (MANDATORY - DO NOT SKIP)\\n"
            "Provide 4-5 specific hydration recommendations:\\n"
            "- Daily water intake target: 8-10 glasses (2.5-3 liters)\\n"
            "- Morning hydration: 2 glasses of water upon waking\\n"
            "- Water intake before meals (20-30 minutes before)\\n"
            "- Herbal teas: Ginger water, Jeera water, Green tea\\n"
            "- Avoiding dehydrating beverages: Excessive caffeine, sugary drinks\\n"
            "\\n"
            "For each, output:\\n"
            "- Category: Hydration\\n"
            "  Recommendation: [specific hydration practice]\\n"
            "  Benefits for Mental Health: [1-2 lines]\\n"
            "  Benefits for Physical Health: [1-2 lines]\\n"
            "\\n"
            "===== CATEGORY 4: LIFESTYLE ===== (MANDATORY - DO NOT SKIP)\\n"
            "Provide 5-6 specific lifestyle recommendations:\\n"
            "- Daily physical activity: 30 minutes of yoga, walking, or exercise\\n"
            "- Screen time management: Limit your screen time\\n"
            "- Stress management: 10-minute meditation, journaling\\n"
            "- Social connections: Quality time with family/friends weekly\\n"
            "- Time with nature: Outdoor walks, sunlight exposure\\n"
            "- Digital detox: Tech-free hours, weekend detox\\n"
            "\\n"
            "For each, output:\\n"
            "- Category: Lifestyle\\n"
            "  Recommendation: [specific lifestyle practice]\\n"
            "  Benefits for Mental Health: [1-2 lines]\\n"
            "  Benefits for Physical Health: [1-2 lines]\\n"
            "\\n"
            "FINAL CHECK BEFORE SUBMISSION:\\n"
            "- Have you included FOOD category? (6-8 items)\\n"
            "- Have you included SLEEPING DISCIPLINE category? (5-6 items)\\n"
            "- Have you included HYDRATION category? (4-5 items)\\n"
            "- Have you included LIFESTYLE category? (5-6 items)\\n"
            "\\n"
            "CRITICAL RULES:\\n"
            "- ALL 4 categories are MANDATORY. Do not skip any.\\n"
            "- Each recommendation MUST have all 4 fields: Category, Recommendation, Benefits for Mental Health, Benefits for Physical Health.\\n"
            "- Be SPECIFIC with examples (not generic).\\n"
            "- Include Indian food context where relevant.\\n"
            "- No markdown tables, no pipes.\\n"
        ),
    }

    prompts = []
    for section in sections:
        blueprint = section_blueprints.get(section, "Include section-specific actionable steps aligned to the heading.")
        
        prompts.append(
            base_prompt +
            f"{blueprint}\n"
            "WRITE ONLY THE FOLLOWING SECTION, using the exact heading text as the first line:\n"
            f"{section}\n"
        )
    
    return prompts

@app.route("/generate-report", methods=["POST"])
def generate_report():
    """
    Flask endpoint for report generation using multi-agent system.
    Maintains backward compatibility with existing frontend.
    """
    try:
        data = request.json
        report_type = data.get("report_type")
        inputs = data.get("inputs")
        
        if not inputs:
            return jsonify({"error": "No inputs provided"}), 400
        
        print(f"\n{'#'*70}")
        print(f"# FLASK ENDPOINT: Report Generation Request Received")
        print(f"# Report Type: {report_type}")
        print(f"# Student: {inputs.get('student_name', 'Unknown')}")
        print(f"{'#'*70}\n")
        
        # Use multi-agent system to generate report
        report_content = generate_report_with_agents(report_type, inputs)
        
        # Generate Word document (UNCHANGED - preserves exact format)
        filename = generate_word_document(report_content, report_type, inputs)
        
        print(f"\n{'#'*70}")
        print(f"# FLASK ENDPOINT: Report Generation Complete")
        print(f"# File: {filename}")
        print(f"{'#'*70}\n")
        
        return jsonify({
            "success": True,
            "content": report_content,
            "filename": filename
        })
        
    except Exception as e:
        print(f"\n[ERROR] Report generation failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        filepath = os.path.join('generated_reports', filename)
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        return send_file(filepath, as_attachment=True, download_name=filename)
    except Exception as e:
        print(f"Error in download_file: {str(e)}")
        return jsonify({'error': str(e)}), 404

# def create_career_report_prompt(inputs):
#     """Create comprehensive prompt for Career Report (Output Fields 1)"""
    
#     prompt = f"""
# Based on the following student profile information, create a comprehensive career-focused education portfolio report:

# **INPUT DATA:**
# - Highest Skills: {', '.join(inputs.get('highest_skills', [])) if inputs.get('highest_skills') else 'N/A'}
# - Thinking Pattern: {inputs.get('thinking_pattern', 'N/A')}
# - Achievement Style: {', '.join(inputs.get('achievement_style', [])) if inputs.get('achievement_style') else 'N/A'}
# - Learning & Communication Style: {', '.join(inputs.get('learning_communication_style', [])) if inputs.get('learning_communication_style') else 'N/A'}
# - Quotients: {', '.join(inputs.get('quotients', [])) if inputs.get('quotients') else 'N/A'}
# - Personality Type: {inputs.get('personality_type', 'N/A')}
# - Suggested Career Roles: {inputs.get('career_roles', 'N/A')}

# # AFTER
# "REQUIRED OUTPUT\n"
# "Please provide a detailed, personalized report with the following sections. "
# "Each section should be comprehensive, around 500-700 words per section, "
# "with detailed explanations, examples, and practical recommendations.\n\n"

# 1. **Suggested Job Roles**
#    - Provide PRIMARY job role (most suitable based on the profile)
#    - Provide SECONDARY job role (alternative career path)
#    - Explain why these roles align with the student's skills, personality, and thinking pattern

# 2. **Industry Specific Requirements**
#    - Detail the specific qualifications, certifications, and skills needed for the suggested roles
#    - Include educational prerequisites and professional requirements
#    - Mention any industry-standard certifications or licenses

# 3. **Emerging Trends and Future Job Prospects**
#    - Analyze current market trends relevant to the suggested careers
#    - Discuss future growth prospects and evolving opportunities
#    - Mention emerging technologies or skills that will be valuable

# 4. **Recommended Internships**
#    - Suggest specific types of internships that would be beneficial
#    - Include ideal companies or organizations to target
#    - Mention expected learning outcomes from these internships

# 5. **Professional Networking and Industry Associations**
#    - List relevant professional organizations and associations to join
#    - Suggest networking strategies and platforms (LinkedIn groups, conferences, etc.)
#    - Mention key industry events or seminars to attend

# 6. **Guidelines for Progress Monitoring & Support**
#    - Provide a framework for tracking career development progress
#    - Suggest mentorship opportunities and support systems
#    - Include key performance indicators (KPIs) for career growth

# 7. **Monitoring Points (Monthly/Quarterly)**
#    - Create specific monthly and quarterly checkpoints
#    - Include measurable goals and milestones
#    - Suggest self-assessment methods and feedback mechanisms

# 8. **Personal Vision Statement**
#    - Craft a personalized vision statement based on the student's profile
#    - Include short-term (1-2 years) and long-term (5-10 years) goals
#    - Make it inspiring and aligned with their strengths

# 9. **Review & Check-In Milestones**
#    - Define specific milestones for 6 months, 1 year, 2 years, and 5 years
#    - Include criteria for evaluating progress at each milestone
#    - Suggest adjustment strategies if goals are not being met

# **FORMAT REQUIREMENTS:**
# - Use clear headings for each section (use ## for main headings)
# - Provide detailed, actionable content
# - Make recommendations specific and personalized
# - Use professional, encouraging tone
# - Include practical examples where relevant
# """
    
#     return prompt

# def create_development_report_prompt(inputs):
#     """Create comprehensive prompt for Development Report (Output Fields 2)"""
    
#     prompt = f"""
# Based on the following student profile information, create a comprehensive personal development and intervention education portfolio report:

# **INPUT DATA:**
# - Highest Skills: {', '.join(inputs.get('highest_skills', [])) if inputs.get('highest_skills') else 'N/A'}
# - Thinking Pattern: {inputs.get('thinking_pattern', 'N/A')}
# - Achievement Style: {', '.join(inputs.get('achievement_style', [])) if inputs.get('achievement_style') else 'N/A'}
# - Learning & Communication Style: {', '.join(inputs.get('learning_communication_style', [])) if inputs.get('learning_communication_style') else 'N/A'}
# - Quotients: {', '.join(inputs.get('quotients', [])) if inputs.get('quotients') else 'N/A'}
# - Personality Type: {inputs.get('personality_type', 'N/A')}
# - Suggested Career Roles: {inputs.get('career_roles', 'N/A')}

# **REQUIRED OUTPUT:**
# Please provide a detailed, personalized development plan with the following sections. Each section should be comprehensive (150-300 words per section):

# 1. **Academic Interventions**
#    - Create a detailed academic development timeline spanning 4-6 years
#    - Include specific courses, certifications, and educational milestones for each year
#    - Break down Year 1 and Year 2 into monthly goals
#    - Provide quarterly goals for Years 3-6
#    - Align academic interventions with the student's learning style and career goals

# 2. **Non-Academic Interventions**
#    - Suggest extracurricular activities, workshops, and skill-building programs
#    - Include leadership development opportunities
#    - Recommend volunteer work or community service aligned with career goals
#    - Suggest creative pursuits and hobbies for holistic development

# 3. **Habit Reengineering**
#    - Identify productive habits to develop based on the student's profile
#    - Suggest habits to eliminate or modify
#    - Provide step-by-step habit formation strategies
#    - Include time management and productivity techniques
#    - Align habits with achievement style and learning patterns

# 4. **Physical Grooming**
#    - Recommend physical fitness routines suitable for the student
#    - Suggest appearance and professional grooming standards for their career path
#    - Include posture, body language, and presence development tips
#    - Recommend sports or physical activities for stress management

# 5. **Psychological Grooming**
#    - Provide mental health and emotional intelligence development strategies
#    - Suggest stress management and resilience-building techniques
#    - Include mindfulness or meditation practices
#    - Recommend ways to develop confidence and self-esteem
#    - Address potential psychological barriers based on personality type

# 6. **Suggested Reading**
#    - **Technical Books:** List 5-7 essential technical books relevant to their career path
#      (Include title, author, and why it's recommended)
#    - **Soft Skills Books:** List 5-7 books for personal development, communication, leadership
#      (Include title, author, and key takeaways)
#    - Create a reading schedule (books per quarter/year)

# 7. **Health Discipline**
#    - Recommend optimal diet plans for cognitive performance and energy
#    - Suggest meal timing and nutritional guidelines
#    - Provide sleep hygiene practices and optimal sleep schedules
#    - Include tips for maintaining consistency
#    - Address stress-eating or irregular eating patterns

# **FORMAT REQUIREMENTS:**
# - Use clear headings for each section (use ## for main headings)
# - Provide detailed, actionable content
# - Make recommendations specific and personalized
# - Use professional, encouraging tone
# - Include practical implementation steps
# - For the 4-6 year plan, use clear timeline markers (Year 1 - Month 1, Month 2, etc.)
# """
    
#     return prompt

def set_cell_background(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_detailed_career_breakdown_table(section_content, doc):
    """
    Parse Section 1: Detailed Career Role Breakdown
    Creates proper Word tables - ONE table per career role with 2 columns (Aspect, Details)
    """
    lines = [l.strip() for l in section_content.split('\n') if l.strip()]
    
    # Extract career role data
    career_roles = []
    current_role = None
    current_data = {}
    
    for line in lines:
        # Detect "Career Role:" lines
        if line.startswith('Career Role:'):
            if current_role and current_data:
                career_roles.append({current_role: current_data})
            
            current_role = line.replace('Career Role:', '').strip()
            current_data = {}
        
        # Parse key-value pairs
        elif ':' in line and current_role:
            key, value = line.split(':', 1)
            key = key.strip()
            value = value.strip()
            current_data[key] = value
    
    # Don't forget last role
    if current_role and current_data:
        career_roles.append({current_role: current_data})
    
    if not career_roles:
        return False
    
    # Create ONE table per career role
    for role_dict in career_roles:
        role_name = list(role_dict.keys())[0]
        role_data = role_dict[role_name]
        
        # Add role as heading
        doc.add_heading(role_name, level=3)
        
        # Create 2-column table for this role
        rows_list = []
        for key, value in role_data.items():
            rows_list.append([key, value])
        
        table = doc.add_table(rows=len(rows_list) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Aspect'
        header_cells[1].text = 'Details'
        
        set_cell_background(header_cells[0], '4472C4')
        set_cell_background(header_cells[1], '4472C4')
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for row_idx, (aspect, details) in enumerate(rows_list):
            row = table.rows[row_idx + 1]
            row.cells[0].text = aspect
            row.cells[1].text = details
            
            # Alternating colors
            if row_idx % 2 == 0:
                set_cell_background(row.cells[0], 'D9E1F2')
                set_cell_background(row.cells[1], 'D9E1F2')
        
        doc.add_paragraph()
    
    return True

def create_industry_requirements_table(section_content, doc):
    """
    Parse Section 2: Industry Specific Requirements
    Output as BULLET POINTS format (not table)
    """
    lines = [l.strip() for l in section_content.split('\n') if l.strip()]
    
    current_role = None
    current_level = None
    
    for line in lines:
        # Detect role heading
        if line.startswith('For ') and ':' in line:
            current_role = line.replace('For ', '').replace(':', '').strip()
            doc.add_heading(f"For {current_role}:", level=3)
            current_level = None
        
        # Detect level heading
        elif 'Level:' in line:
            level_name = line.split(':')[0].strip()
            doc.add_heading(level_name, level=4)
            current_level = level_name
        
        # Add content as bullet points
        elif line.startswith('Certification Name:') or line.startswith('Application Process:') or \
             line.startswith('Duration:') or line.startswith('Assistance Resources:'):
            # Format: Label: Value
            if ':' in line:
                label, value = line.split(':', 1)
                doc.add_paragraph(f"{label.strip()}: {value.strip()}", style='List Bullet')
        
        elif line.startswith('-'):
            # Already a bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
    
    doc.add_paragraph()
    return True


def create_emerging_trends_table(section_content, doc):
    """
    Parse Section 3: Emerging Trends with Past | Present | Future
    Enhanced to handle markdown table format from AI
    """
    lines = [l.strip() for l in section_content.split('\n') if l.strip()]
    
    current_role = None
    role_tables = {}
    
    for line in lines:
        # Detect role headings
        if line and not line.startswith('|') and not line.startswith('-') and len(line) < 50:
            # Likely a role heading
            current_role = line.strip()
            role_tables[current_role] = []
        elif current_role and line.startswith('|') and '|' in line[1:]:
            # This is a table row
            role_tables[current_role].append(line)
    
    if not role_tables:
        return False
    
    for role, table_lines in role_tables.items():
        doc.add_heading(role, level=3)
        
        # Parse markdown table
        parsed_rows = []
        for line in table_lines:
            if '---' in line:  # Skip separator line
                continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                parsed_rows.append(cells)
        
        if len(parsed_rows) < 2:  # Need at least header + 1 data row
            continue
        
        # Create Word table
        table = doc.add_table(rows=len(parsed_rows), cols=len(parsed_rows[0]))
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for row_idx, row_data in enumerate(parsed_rows):
            row = table.rows[row_idx]
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = cell_data
            
            # Format header row
            if row_idx == 0:
                for cell in row.cells:
                    set_cell_background(cell, '4472C4')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
            # Alternating colors for data rows
            elif row_idx % 2 == 1:
                for cell in row.cells:
                    set_cell_background(cell, 'D9E1F2')
        
        doc.add_paragraph()
    
    return True

def create_internships_table(section_content, doc):
    """
    Parse Section 4: Recommended Internships
    Format: Internship Type | Industries (Small/Medium/Large as bullets in ONE cell) | Expected Outcomes
    FIXED: Skip header rows and separator rows from markdown table
    """
    lines = [l.strip() for l in section_content.split('\n') if l.strip()]
    
    current_role = None
    internship_data = {}
    
    for line in lines:
        # Detect role heading
        if line.startswith('For ') and ':' in line:
            current_role = line.replace('For ', '').replace(':', '').strip()
            internship_data[current_role] = []
        
        # SKIP header rows and separator rows (CRITICAL FIX)
        elif line.startswith('|') and current_role and '|' in line[1:]:
            cells = [c.strip() for c in line.split('|') if c.strip()]
            
            if not cells:
                continue
            
            # Skip separator rows (all dashes)
            if all(c in '-' for c in ''.join(cells)):
                continue
            
            # Skip if it matches header text exactly
            if 'Internship Type' in cells[0] or 'Industries' in cells[0]:
                continue
            
            # Process as data row only
            if len(cells) >= 3:
                internship_type = cells[0]
                industries = cells[1] if len(cells) > 1 else ""
                outcomes = cells[2] if len(cells) > 2 else ""
                
                # Double-check this is NOT a header
                if internship_type not in ['Internship Type', 'Industries', 'Expected Outcomes']:
                    internship_data[current_role].append({
                        'type': internship_type,
                        'industries': industries,
                        'outcomes': outcomes
                    })
    
    if not internship_data:
        return False
    
    # Create tables for each role
    for role, internships in internship_data.items():
        doc.add_heading(f"For {role}:", level=3)
        
        if not internships:
            continue
        
        # Create table: Internship Type | Industries | Expected Outcomes
        table = doc.add_table(rows=len(internships) + 1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row - ONLY add headers once
        headers = ['Internship Type', 'Industries', 'Expected Outcomes']
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            set_cell_background(cell, '4472C4')
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for row_idx, internship in enumerate(internships):
            row = table.rows[row_idx + 1]
            row.cells[0].text = internship['type']
            
            # Industries cell - parse into bullet points
            industries_cell = row.cells[1]
            industries_cell.text = ""
            
            industries_text = internship['industries']
            
            small_part = ""
            medium_part = ""
            large_part = ""
            
            if "Small:" in industries_text:
                small_start = industries_text.index("Small:") + 6
                small_end = industries_text.index("Medium:") if "Medium:" in industries_text else len(industries_text)
                small_part = industries_text[small_start:small_end].strip()
            
            if "Medium:" in industries_text:
                medium_start = industries_text.index("Medium:") + 7
                medium_end = industries_text.index("Large:") if "Large:" in industries_text else len(industries_text)
                medium_part = industries_text[medium_start:medium_end].strip()
            
            if "Large:" in industries_text:
                large_start = industries_text.index("Large:") + 6
                large_part = industries_text[large_start:].strip()
            
            # Add as bullets
            if small_part:
                p = industries_cell.paragraphs[0] if industries_cell.paragraphs else industries_cell.add_paragraph()
                p.text = f"• Small: {small_part}"
            
            if medium_part:
                p = industries_cell.add_paragraph()
                p.text = f"• Medium: {medium_part}"
            
            if large_part:
                p = industries_cell.add_paragraph()
                p.text = f"• Large: {large_part}"
            
            row.cells[2].text = internship['outcomes']
            
            # Alternating colors
            if row_idx % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, 'D9E1F2')
        
        doc.add_paragraph()
    
    # Add application pipeline advice
    pipeline_idx = next((i for i, line in enumerate(lines) if 'Application Pipeline' in line), None)
    if pipeline_idx:
        doc.add_heading("Application Pipeline Advice:", level=3)
        for line in lines[pipeline_idx:]:
            if line.startswith('-'):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif not line.startswith('For ') and not line.startswith('|') and line:
                doc.add_paragraph(line)
    
    return True

def create_networking_table(section_content, doc):
    """
    Parse Section 5: Professional Networking
    Fixed to handle structured format with bullet points
    """
    lines = [l.strip() for l in section_content.split('\n') if l.strip()]
    
    # Extract role-based data
    role_tables = {}
    current_role = None
    current_section = None
    current_items = []
    
    for line in lines:
        # Skip empty lines
        if not line:
            continue
        
        # Detect role headings (e.g., "For Data Scientist:")
        if line.startswith('For ') and ':' in line:
            # Save previous role if exists
            if current_role:
                if current_role not in role_tables:
                    role_tables[current_role] = {'associations': [], 'events': [], 'strategy': []}
                if current_section:
                    role_tables[current_role][current_section] = current_items
            
            current_role = line.replace('For ', '').replace(':', '').strip()
            current_section = None
            current_items = []
        
        # Detect section headings within role
        elif 'Professional Associations:' in line:
            if current_role and current_section:
                if current_role not in role_tables:
                    role_tables[current_role] = {'associations': [], 'events': [], 'strategy': []}
                role_tables[current_role][current_section] = current_items
            current_section = 'associations'
            current_items = []
        
        elif 'Industry Events:' in line:
            if current_role and current_section:
                if current_role not in role_tables:
                    role_tables[current_role] = {'associations': [], 'events': [], 'strategy': []}
                role_tables[current_role][current_section] = current_items
            current_section = 'events'
            current_items = []
        
        elif 'Networking Strategy:' in line:
            if current_role and current_section:
                if current_role not in role_tables:
                    role_tables[current_role] = {'associations': [], 'events': [], 'strategy': []}
                role_tables[current_role][current_section] = current_items
            current_section = 'strategy'
            current_items = []
        
        # Collect bullet items
        elif line.startswith('- ') and current_section:
            current_items.append(line[2:])  # Remove "- " prefix
    
    # Don't forget last role and section
    if current_role and current_section:
        if current_role not in role_tables:
            role_tables[current_role] = {'associations': [], 'events': [], 'strategy': []}
        role_tables[current_role][current_section] = current_items
    
    if not role_tables:
        return False
    
    # Create separate table for each role
    for role, data in role_tables.items():
        # Add role heading
        doc.add_heading(f"For {role}:", level=3)
        
        # Create table with 3 columns
        table = doc.add_table(rows=2, cols=3)  # Header + 1 data row
        table.style = 'Light Grid Accent 1'
        
        # Header row
        headers = ['Professional Associations', 'Industry Events', 'Networking Strategy']
        header_cells = table.rows[0].cells
        
        for col_idx, header in enumerate(headers):
            header_cells[col_idx].text = header
            set_cell_background(header_cells[col_idx], '4472C4')
            for paragraph in header_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data row
        row = table.rows[1]
        
        # Associations cell - with bullets
        assoc_cell = row.cells[0]
        assoc_cell.text = ""
        for idx, item in enumerate(data.get('associations', [])):
            if idx == 0:
                p = assoc_cell.paragraphs[0]
            else:
                p = assoc_cell.add_paragraph()
            p.text = f"• {item}"
        
        # Events cell - with bullets
        events_cell = row.cells[1]
        events_cell.text = ""
        for idx, item in enumerate(data.get('events', [])):
            if idx == 0:
                p = events_cell.paragraphs[0]
            else:
                p = events_cell.add_paragraph()
            p.text = f"• {item}"
        
        # Strategy cell - with bullets
        strategy_cell = row.cells[2]
        strategy_cell.text = ""
        for idx, item in enumerate(data.get('strategy', [])):
            if idx == 0:
                p = strategy_cell.paragraphs[0]
            else:
                p = strategy_cell.add_paragraph()
            p.text = f"• {item}"
        
        # Alternating colors
        set_cell_background(row.cells[0], 'D9E1F2')
        set_cell_background(row.cells[1], 'D9E1F2')
        set_cell_background(row.cells[2], 'D9E1F2')
        
        doc.add_paragraph()
    
    return True

def create_progress_monitoring_table(section_content, doc):
    """
    Parse Section 6: Progress Monitoring
    Format: Horizontal table | Aspect | Role1 | Role2 | Role3 |
    """
    lines = [l.strip() for l in section_content.split('\n') if l.strip()]
    
    # Find table lines
    table_lines = [line for line in lines if line.startswith('|')]
    
    if not table_lines:
        return False
    
    # Parse markdown table
    parsed_rows = []
    for line in table_lines:
        if '---' in line or all(c in '|-: ' for c in line):
            continue
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            parsed_rows.append(cells)
    
    if len(parsed_rows) < 2:
        return False
    
    # Create Word table
    table = doc.add_table(rows=len(parsed_rows), cols=len(parsed_rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for row_idx, row_data in enumerate(parsed_rows):
        row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row.cells):
                row.cells[col_idx].text = cell_data
        
        # Header formatting (both header row and first column)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_background(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            # First column (Aspect) also gets bold
            row.cells[0].paragraphs[0].runs[0].font.bold = True if row.cells[0].paragraphs[0].runs else None
            
            # Alternating colors for data cells only
            if row_idx % 2 == 1:
                for cell in row.cells[1:]:
                    set_cell_background(cell, 'D9E1F2')
    
    doc.add_paragraph()
    return True

def _parse_year_block(lines):
    """
    Helper: parse structured year-wise blocks
    FIXED: Normalize field names to handle both 'Action plan:' and 'Action Plan:'
    """
    year_data = {}
    current_year = None
    current_row = None

    for line in lines:
        # Skip empty lines
        if not line or not line.strip():
            continue

        stripped = line.strip()

        # Detect year headers: "Year 1:", "Year 2:", etc.
        if stripped.startswith("Year ") and stripped.endswith(":"):
            current_year = stripped.rstrip(":")
            if current_year not in year_data:
                year_data[current_year] = []
            current_row = None
            continue

        # Detect start of new row - look for "Month:"
        if "Month:" in stripped:
            # Save previous row if exists
            if current_year and current_row:
                year_data[current_year].append(current_row)
            
            if current_year is None:
                continue

            # Start new row
            current_row = {}
            
            # Extract Month value
            month_value = stripped.replace("- Month:", "").replace("Month:", "").strip()
            current_row["Month"] = month_value
            continue

        # Parse other fields (Activity, Technical Skills, Soft Skills, etc.)
        if current_year and current_row and ":" in stripped:
            # Split key and value
            key, value = stripped.split(":", 1)
            key = key.strip().lstrip("-").strip()  # Remove leading dash if present
            value = value.strip()

            # NORMALIZE KEY NAMES - Handle both capitalized and lowercase variations
            # Convert "Action plan:" to "Action Plan:"
            if key.lower() == "action plan":
                key = "Action Plan"
            elif key.lower() == "learning outcomes":
                key = "Learning Outcomes"
            elif key.lower() == "learning outcome":
                key = "Learning Outcome"
            elif key.lower() == "habits to develop":
                key = "Habits to Develop"
            elif key.lower() == "soft skills":
                key = "Soft Skills"
            elif key.lower() == "technical skills":
                key = "Technical Skills"
            elif key.lower() == "learning material":
                key = "Learning Material"
            elif key.lower() == "psychological skills developed":
                key = "Psychological Skills Developed"
            elif key.lower() == "physical & mental skills developed":
                key = "Physical & Mental Skills Developed"

            # List of valid field names (normalized)
            valid_fields = [
                "Activity", "Action Plan", "Technical Skills", "Soft Skills", "Learning Material", "Objective",
                "Learning Outcome", "Habits to Develop", "Learning Outcomes",
                "Physical & Mental Skills Developed", "Psychological Skills Developed"
            ]

            if key in valid_fields and value:
                current_row[key] = value

    # Append last row if exists
    if current_year and current_row:
        year_data[current_year].append(current_row)

    return year_data

def create_academic_interventions_tables(section_content, doc):
    """
    Section 1 - Academic Interventions
    Columns: Month, Activity, Technical Skills, Soft Skills, Learning Material, Objective
    Year-wise tables (Year 1, Year 2, Year 3)
    """
    lines = [l.strip() for l in section_content.split("\n") if l.strip()]
    year_data = _parse_year_block(lines)
    if not year_data:
        return False

    columns = ["Month", "Activity", "Technical Skills", "Soft Skills", "Learning Material", "Objective"]

    for year in ["Year 1", "Year 2", "Year 3"]:
        if year not in year_data or not year_data[year]:
            continue

        doc.add_heading(year, level=3)
        table = doc.add_table(rows=len(year_data[year]) + 1, cols=len(columns))
        table.style = "Light Grid Accent 1"

        # Header
        header_cells = table.rows[0].cells
        for idx, col in enumerate(columns):
            header_cells[idx].text = col
            set_cell_background(header_cells[idx], "4472C4")
            for p in header_cells[idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        # Rows
        for row_idx, row_data in enumerate(year_data[year]):
            row = table.rows[row_idx + 1]
            for col_idx, col in enumerate(columns):
                # Get value, handle missing gracefully
                value = row_data.get(col, "")
                # Remove extra whitespace and ensure clean text
                value = value.strip() if value else ""
                row.cells[col_idx].text = value
                # Set text wrapping
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row_idx % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, "D9E1F2")

        doc.add_paragraph()

    return True


def create_non_academic_interventions_tables(section_content, doc):
    """
    Section 2 - Non-Academic Interventions
    Columns: Month, Activity, Technical Skills, Soft Skills, Learning Outcome, Objective
    """
    lines = [l.strip() for l in section_content.split("\n") if l.strip()]
    year_data = _parse_year_block(lines)
    if not year_data:
        return False

    columns = ["Month", "Activity", "Technical Skills", "Soft Skills", "Learning Outcome", "Objective"]

    for year in ["Year 1", "Year 2", "Year 3"]:
        if year not in year_data or not year_data[year]:
            continue

        doc.add_heading(year, level=3)
        table = doc.add_table(rows=len(year_data[year]) + 1, cols=len(columns))
        table.style = "Light Grid Accent 1"

        header_cells = table.rows[0].cells
        for idx, col in enumerate(columns):
            header_cells[idx].text = col
            set_cell_background(header_cells[idx], "4472C4")
            for p in header_cells[idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, row_data in enumerate(year_data[year]):
            row = table.rows[row_idx + 1]
            for col_idx, col in enumerate(columns):
                # Get value, handle missing gracefully
                value = row_data.get(col, "")
                # Remove extra whitespace and ensure clean text
                value = value.strip() if value else ""
                row.cells[col_idx].text = value
                # Set text wrapping
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row_idx % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, "D9E1F2")

        doc.add_paragraph()

    return True


def create_habit_reengineering_tables(section_content, doc):
    """
    Section 3 - Habit Reengineering
    Columns: Month, Activity, Action Plan, Objective, Habits to Develop, Soft Skills, Learning Outcomes
    
    ACTION PLAN: Provides detailed explanation of HOW to perform the activity
    """
    lines = [l.strip() for l in section_content.split("\n") if l.strip()]
    year_data = _parse_year_block(lines)
    
    if not year_data:
        return False
    
    # UPDATED: Added "Action Plan" column after "Activity"
    columns = ["Month", "Activity", "Action Plan", "Objective", "Habits to Develop", "Soft Skills", "Learning Outcomes"]
    
    for year in ["Year 1", "Year 2", "Year 3"]:
        if year not in year_data or not year_data[year]:
            continue
        
        doc.add_heading(year, level=3)
        
        table = doc.add_table(rows=len(year_data[year]) + 1, cols=len(columns))
        table.style = "Light Grid Accent 1"
        
        # Header row
        header_cells = table.rows[0].cells
        for idx, col in enumerate(columns):
            header_cells[idx].text = col
            set_cell_background(header_cells[idx], "4472C4")
            
            for p in header_cells[idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for row_idx, row_data in enumerate(year_data[year]):
            row = table.rows[row_idx + 1]
            
            for col_idx, col in enumerate(columns):
                # Get value, handle missing gracefully
                value = row_data.get(col, "")
                # Remove extra whitespace and ensure clean text
                value = value.strip() if value else ""
                row.cells[col_idx].text = value
                
                # Set text wrapping
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Alternating row colors
            if row_idx % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, "D9E1F2")
        
        doc.add_paragraph()
    
    return True


def create_physical_grooming_tables(section_content, doc):
    """
    Section 4 - Physical Grooming
    Columns: Month, Activity, Objective, Physical & Mental Skills Developed, Soft Skills, Learning Outcomes
    """
    lines = [l.strip() for l in section_content.split("\n") if l.strip()]
    year_data = _parse_year_block(lines)
    if not year_data:
        return False

    columns = ["Month", "Activity", "Objective", "Physical & Mental Skills Developed", "Soft Skills", "Learning Outcomes"]

    for year in ["Year 1", "Year 2", "Year 3"]:
        if year not in year_data or not year_data[year]:
            continue

        doc.add_heading(year, level=3)
        table = doc.add_table(rows=len(year_data[year]) + 1, cols=len(columns))
        table.style = "Light Grid Accent 1"

        header_cells = table.rows[0].cells
        for idx, col in enumerate(columns):
            header_cells[idx].text = col
            set_cell_background(header_cells[idx], "4472C4")
            for p in header_cells[idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, row_data in enumerate(year_data[year]):
            row = table.rows[row_idx + 1]
            for col_idx, col in enumerate(columns):
                # Get value, handle missing gracefully
                value = row_data.get(col, "")
                # Remove extra whitespace and ensure clean text
                value = value.strip() if value else ""
                row.cells[col_idx].text = value
                # Set text wrapping
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            if row_idx % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, "D9E1F2")

        doc.add_paragraph()

    return True


def create_psychological_grooming_tables(section_content, doc):
    """
    Section 5 - Psychological Grooming
    Columns: Month, Activity, Objective, Psychological Skills Developed, Soft Skills, Learning Outcomes
    """
    lines = [l.strip() for l in section_content.split("\n") if l.strip()]
    year_data = _parse_year_block(lines)
    if not year_data:
        return False

    columns = ["Month", "Activity", "Objective", "Psychological Skills Developed", "Soft Skills", "Learning Outcomes"]

    for year in ["Year 1", "Year 2", "Year 3"]:
        if year not in year_data or not year_data[year]:
            continue

        doc.add_heading(year, level=3)
        table = doc.add_table(rows=len(year_data[year]) + 1, cols=len(columns))
        table.style = "Light Grid Accent 1"

        header_cells = table.rows[0].cells
        for idx, col in enumerate(columns):
            header_cells[idx].text = col
            set_cell_background(header_cells[idx], "4472C4")
            for p in header_cells[idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, row_data in enumerate(year_data[year]):
            row = table.rows[row_idx + 1]
            for col_idx, col in enumerate(columns):
                # Get value, handle missing gracefully
                value = row_data.get(col, "")
                # Remove extra whitespace and ensure clean text
                value = value.strip() if value else ""
                row.cells[col_idx].text = value
                # Set text wrapping
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            if row_idx % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, "D9E1F2")

        doc.add_paragraph()

    return True


def create_suggested_reading_table(section_content, doc):
    """
    Section 6 - Suggested Reading
    Columns: Book Name, Author, Publication, Why Should This Book Be Read?
    """
    lines = [l.strip() for l in section_content.split("\n") if l.strip()]
    rows = []
    current = {}

    for line in lines:
        if line.startswith("- Book Name:"):
            if current:
                rows.append(current)
            current = {"Book Name": line.split(":", 1)[1].strip()}
        elif ":" in line and current and not line.startswith("Year "):
            key, value = line.split(":", 1)
            current[key.strip()] = value.strip()

    if current:
        rows.append(current)

    if not rows:
        return False

    columns = ["Book Name", "Author", "Publication", "Why Should This Book Be Read?"]
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    for idx, col in enumerate(columns):
        header_cells[idx].text = col
        set_cell_background(header_cells[idx], "4472C4")
        for p in header_cells[idx].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, col in enumerate(columns):
            row.cells[col_idx].text = row_data.get(col, "")
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "D9E1F2")

    doc.add_paragraph()
    return True


def create_diet_sleep_table(section_content, doc):
    """
    Section 7 - Health Discipline
    Columns: Category, Recommendation, Benefits for Mental Health, Benefits for Physical Health
    """
    lines = [l.strip() for l in section_content.split("\n") if l.strip()]
    rows = []
    current = {}

    for line in lines:
        if line.startswith("- Category:"):
            if current:
                rows.append(current)
            current = {"Category": line.split(":", 1)[1].strip()}
        elif ":" in line and current:
            key, value = line.split(":", 1)
            current[key.strip()] = value.strip()

    if current:
        rows.append(current)

    if not rows:
        return False

    columns = ["Category", "Recommendation", "Benefits for Mental Health", "Benefits for Physical Health"]
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    for idx, col in enumerate(columns):
        header_cells[idx].text = col
        set_cell_background(header_cells[idx], "4472C4")
        for p in header_cells[idx].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, col in enumerate(columns):
            row.cells[col_idx].text = row_data.get(col, "")
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "D9E1F2")

    doc.add_paragraph()
    return True

def parse_section_to_table(section_heading, section_content, doc, report_type):
    """
    Parse bullet point section content and convert to structured table format
    """
    
    # Define which sections should be converted to tables
    career_table_sections = [
        "1. Detailed Career Role Breakdown",
        "2. Industry Specific Requirements",
        "3. Emerging Trends and Future Job Prospects",
        "4. Recommended Internships",
        "5. Professional Networking and Industry Associations",
        "6. Guidelines for Progress Monitoring & Support"
    ]
    
    development_table_sections = [
        "1. Academic Interventions",
        "2. Non-Academic Interventions",
        "3. Habit Reengineering",
        "4. Physical Grooming",
        "5. Psychological Grooming",
        "6. Suggested Reading",
        "7. Health Discipline"
    ]
    
    # Check if this section should be converted to table
    sections_to_convert = career_table_sections if report_type == 'career' else development_table_sections
    
    if not any(section_heading.startswith(s.split('.')[0]) for s in sections_to_convert):
        return None  # Don't convert this section
    
    # Add section heading
    doc.add_heading(section_heading, 2)

    # Handle specific sections with custom table formats
    if "1. Detailed Career Role Breakdown" in section_heading:
        return create_detailed_career_breakdown_table(section_content, doc)
    elif "2. Industry Specific Requirements" in section_heading:
        return create_industry_requirements_table(section_content, doc)
    elif "3. Emerging Trends" in section_heading:
        return create_emerging_trends_table(section_content, doc)
    elif "4. Recommended Internships" in section_heading:
        return create_internships_table(section_content, doc)
    elif "5. Professional Networking" in section_heading:
        return create_networking_table(section_content, doc)
    elif "6. Guidelines for Progress Monitoring" in section_heading:
        return create_progress_monitoring_table(section_content, doc)
    
        # DEVELOPMENT SECTIONS (tables)
    if report_type == 'development':
        if "1. Academic Interventions" in section_heading:
            return create_academic_interventions_tables(section_content, doc)
        elif "2. Non-Academic Interventions" in section_heading:
            return create_non_academic_interventions_tables(section_content, doc)
        elif "3. Habit Reengineering" in section_heading:
            return create_habit_reengineering_tables(section_content, doc)
        elif "4. Physical Grooming" in section_heading:
            return create_physical_grooming_tables(section_content, doc)
        elif "5. Psychological Grooming" in section_heading:
            return create_psychological_grooming_tables(section_content, doc)
        elif "6. Suggested Reading" in section_heading:
            return create_suggested_reading_table(section_content, doc)
        elif "7. Health Discipline" in section_heading:
            return create_diet_sleep_table(section_content, doc)

    # Parse content by roles (split by "For <Role>:" or role-specific patterns)
    lines = [l.strip() for l in section_content.split('\n') if l.strip()]
    
    # Detect roles mentioned in content
    roles = []
    role_data = {}
    current_role = None
    
    for line in lines:
        # Check if line indicates a new role section
        if line.startswith('For ') and ':' in line:
            current_role = line.split(':')[0].replace('For ', '').strip()
            roles.append(current_role)
            role_data[current_role] = []
        elif line.startswith('- ') or line.startswith('• '):
            if current_role:
                role_data[current_role].append(line[2:].strip())
            else:
                # General content (not role-specific)
                if 'General' not in role_data:
                    role_data['General'] = []
                role_data['General'].append(line[2:].strip())
    
    # If no roles detected, create a simple list table
    if not roles or len(role_data) == 0:
        # Fallback: Create 2-column table (Category | Details)
        bullets = [l[2:].strip() for l in lines if l.startswith('- ') or l.startswith('• ')]
        
        if bullets:
            table = doc.add_table(rows=len(bullets) + 1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Category'
            header_cells[1].text = 'Details'
            
            # Format header
            for cell in header_cells:
                set_cell_background(cell, '4472C4')  # Blue header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Add data rows with alternating colors
            for idx, bullet in enumerate(bullets):
                row = table.rows[idx + 1]
                # Try to split by first colon or dash
                if ':' in bullet:
                    parts = bullet.split(':', 1)
                    row.cells[0].text = parts[0].strip()
                    row.cells[1].text = parts[1].strip()
                else:
                    row.cells[0].text = f"Point {idx + 1}"
                    row.cells[1].text = bullet
                
                # Alternating row colors
                if idx % 2 == 0:
                    for cell in row.cells:
                        set_cell_background(cell, 'D9E1F2')  # Light blue
        
        doc.add_paragraph()  # Spacing
        return True
    
    # Create role-based table
    num_roles = len(roles)
    table = doc.add_table(rows=1, cols=num_roles + 1)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Aspect'
    for idx, role in enumerate(roles):
        header_cells[idx + 1].text = role
    
    # Format header with blue background
    for cell in header_cells:
        set_cell_background(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Find all unique aspects (merge data from all roles)
    all_aspects = set()
    for role, data in role_data.items():
        for item in data:
            # Extract aspect name (text before colon if exists)
            if ':' in item:
                aspect = item.split(':')[0].strip()
                all_aspects.add(aspect)
    
    # If no aspects, just list items
    if not all_aspects:
        max_items = max(len(data) for data in role_data.values())
        for i in range(max_items):
            row = table.add_row()
            row.cells[0].text = f"Point {i + 1}"
            for idx, role in enumerate(roles):
                if i < len(role_data.get(role, [])):
                    row.cells[idx + 1].text = role_data[role][i]
            
            # Alternating colors
            if i % 2 == 0:
                for cell in row.cells:
                    set_cell_background(cell, 'D9E1F2')
    else:
        # Organize by aspects
        row_idx = 0
        for aspect in sorted(all_aspects):
            row = table.add_row()
            row.cells[0].text = aspect
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            for idx, role in enumerate(roles):
                matching_items = [item for item in role_data.get(role, []) if item.startswith(aspect + ':')]
                if matching_items:
                    content = matching_items[0].split(':', 1)[1].strip()
                    row.cells[idx + 1].text = content
                else:
                    row.cells[idx + 1].text = 'N/A'
            
            # Alternating colors
            if row_idx % 2 == 0:
                for cell in row.cells[1:]:  # Skip first column (aspect name)
                    set_cell_background(cell, 'D9E1F2')
            row_idx += 1
    
    doc.add_paragraph()  # Spacing
    return True

def generate_word_document(content, report_type, inputs):
    """Generate a formatted Word document from the report content"""
    
    # Create reports directory if it doesn't exist
    os.makedirs('generated_reports', exist_ok=True)
    
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title_text = "Career Development Report" if report_type == 'career' else "Personal Development & Intervention Report"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Add input summary section
    doc.add_heading('Student Profile Summary', 1)
    
    input_table = doc.add_table(rows=10, cols=2)
    input_table.style = 'Light Grid Accent 1'
    
    input_data = [
        ('Name of Student', inputs.get('sname', 'NA')),
        ('Standard / Year', inputs.get('standard', 'NA')),
        ('Board', inputs.get('board', 'NA')),
        ('Highest Skills', format_skills_with_percentages(inputs.get('highest_skills', []), inputs.get('skillpercentages', {}))),
        ('Thinking Pattern', inputs.get('thinking_pattern', 'N/A')),
        ('Achievement Style', format_skills_with_percentages(inputs.get('achievement_style', []), inputs.get('achievementpercentages', {}))),
        ('Learning & Communication Style', format_skills_with_percentages(inputs.get('learning_communication_style', []), inputs.get('learningpercentages', {}))),
        ('Quotients', format_skills_with_percentages(inputs.get('quotients', []), inputs.get('quotientpercentages', {}))),
        ('Personality Type', inputs.get('personality_type', 'N/A')),
        ('Suggested Career Roles', inputs.get('career_roles', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(input_data):
        row = input_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Bold the label
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Add the generated content
    doc.add_heading('Detailed Report', 1)

    # Sanitize content: remove emojis and asterisks
    def remove_emojis_and_asterisks(text):
        # Remove emojis by filtering out characters in the 'So' (symbol, other) or 'Cs' (surrogate) categories
        cleaned = "".join(
            ch for ch in text
            if unicodedata.category(ch) not in ("So", "Cs")
        )
        # Remove asterisks used as decoration
        cleaned = cleaned.replace("*", "")
        return cleaned

    content = remove_emojis_and_asterisks(content)

    # Parse and format the content - Split into sections for table conversion
    sections = content.split('\n\n')  # Split by double newline (sections)
    current_section_heading = None
    current_section_content = []

    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        lines = section.split('\n')
        first_line = lines[0].strip()
        
        # Check if this is a section heading (numbered sections like "1. Suggested Job Roles")
        if re.match(r'^\d+\.\s+[A-Z]', first_line):
            # Process previous section if exists
            if current_section_heading and current_section_content:
                section_text = '\n'.join(current_section_content)
                converted = parse_section_to_table(current_section_heading, section_text, doc, report_type)
                if not converted:
                    # Fallback to regular formatting
                    doc.add_heading(current_section_heading, 2)
                    for line in current_section_content:
                        line = line.strip()
                        if line.startswith('-') or line.startswith('•'):
                            doc.add_paragraph(line[2:].strip(), style='List Bullet')
                        else:
                            doc.add_paragraph(line)
                    doc.add_paragraph()
            
            # Start new section
            current_section_heading = first_line
            current_section_content = lines[1:]  # Rest of the section
        else:
            # Continue current section
            if current_section_heading:
                current_section_content.extend(lines)
            else:
                # No section yet, process as regular content
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    if line.startswith('##'):
                        heading_text = line.replace('#', '').strip()
                        doc.add_heading(heading_text, 2)
                    elif line.startswith('**') and line.endswith('**'):
                        heading_text = line.replace('**', '').strip()
                        doc.add_heading(heading_text, 2)
                    elif line.startswith('###'):
                        heading_text = line.replace('#', '').strip()
                        doc.add_heading(heading_text, 3)
                    elif line.startswith('-') or line.startswith('•'):
                        doc.add_paragraph(line[1:].strip(), style='List Bullet')
                    elif line.startswith(tuple(str(i) + '.' for i in range(1, 10))):
                        doc.add_paragraph(line[2:].strip(), style='List Number')
                    else:
                        doc.add_paragraph(line)

    # Process last section
    if current_section_heading and current_section_content:
        section_text = '\n'.join(current_section_content)
        converted = parse_section_to_table(current_section_heading, section_text, doc, report_type)
        if not converted:
            doc.add_heading(current_section_heading, 2)
            for line in current_section_content:
                line = line.strip()
                if line.startswith('-') or line.startswith('•'):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line)
            doc.add_paragraph()
    
    # Generate filename
    # timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    student_name = inputs.get('sname', 'N/A')
    student_name_clean = sanitize_filename(str(student_name))

    # Career roles
    career_roles = inputs.get('career_roles', 'N/A')
    career_roles_clean = sanitize_filename(str(career_roles))
    report_name = 'Career_Report' if report_type == 'career' else 'Development_Report'
    filename = f"{report_name}_{student_name_clean}_{career_roles_clean}.docx"
    filepath = os.path.join('generated_reports', filename)
    
    # Save document
    doc.save(filepath)
    
    return filename

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
